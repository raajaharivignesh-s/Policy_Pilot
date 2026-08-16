import docx

doc = docx.Document("../knowledge_base/raw/education/Education Schemes.docx")
text = []
for para in doc.paragraphs:
    text.append(para.text)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            text.append(cell.text)

full_text = "\n".join(text)

print("Found occurrences of 'pudh':")
for line in full_text.split("\n"):
    if "pudh" in line.lower() or "pudhavan" in line.lower() or "pudhalvan" in line.lower():
        safe_line = line.encode('ascii', errors='ignore').decode('ascii')
        print("MATCH:", safe_line)
